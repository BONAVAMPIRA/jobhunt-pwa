#!/usr/bin/env python3
"""
docx_render.py — Rendu Markdown -> .docx fidele au format CV TI Quebec.

Remplace la conversion generique de l'ancien docmaker, qui ne gerait que **gras**
(les *italiques* — dates, contexte — sortaient avec leurs asterisques) et n'avait
ni dates alignees a droite, ni pied de page/pagination.

Sources de format : cv/reference/format-cv-ti.md + cv/bases/*.docx (Arial, marges
1,9 cm, nom ~16pt, sections 12pt souligne, outils en gras).

API :
    render(md_text, out_path, kind)   kind in {cv, lm, salaire, guide}
CLI :
    python docx_render.py <input.md> <output.docx> <kind>
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Arial"

# Ligne de dates : entierement en italique ET contient une annee ou un mois FR.
_MONTHS = (r"janvier|f[ée]vrier|mars|avril|mai|juin|juillet|ao[uû]t|"
           r"septembre|octobre|novembre|d[ée]cembre|pr[ée]sent|aujourd")
_DATE_RE = re.compile(rf"(\b(19|20)\d{{2}}\b|{_MONTHS})", re.IGNORECASE)
_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def _add_inline(paragraph, text, size, base_italic=False):
    """Ajoute des runs en gerant **gras** et *italique* imbriques."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        run = None
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1]); run.italic = True
        else:
            run = paragraph.add_run(part)
        if base_italic:
            run.italic = True
        run.font.name = FONT
        run.font.size = Pt(size)


def _section_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "2E74B5")):
        bottom.set(qn(k), v)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _field(paragraph, code):
    """Insere un champ Word (PAGE / NUMPAGES) qui se met a jour a l'ouverture."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = code
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)
    run.font.name = FONT; run.font.size = Pt(8)


def _footer(section, left_text="Références fournies sur demande"):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(left_text + "\t\tPage "); r.font.name = FONT; r.font.size = Pt(8)
    _field(p, "PAGE")
    r2 = p.add_run(" / "); r2.font.name = FONT; r2.font.size = Pt(8)
    _field(p, "NUMPAGES")


def _new_doc(margin_cm=1.9):
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = 12240, 15840  # US Letter (twips)
    m = Inches(round(margin_cm / 2.54, 3))
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = m
    return doc, s


def _heading(doc, text, size, bold=True, border=False, upper=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text.upper() if upper else text)
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold
    if border:
        _section_border(p)
    return p


def render_cv(md_text, out_path):
    """CV au format TI Quebec : nom, titre, sections soulignees, dates a droite, pied."""
    doc, section = _new_doc(margin_cm=1.9)
    _footer(section)
    header_done = False  # nom/titre/PRIIME/contact traites en haut

    lines = md_text.split("\n")
    for raw in lines:
        s = raw.strip()
        if not s or s.startswith("---"):
            continue

        if s.startswith("# "):                       # Nom
            _heading(doc, s[2:].strip(), 16, bold=True)
            header_done = False
            continue
        if s.startswith("## "):                       # Section
            _heading(doc, s[3:].strip(), 12, bold=True, border=True, upper=True)
            header_done = True
            continue
        if s.startswith("### "):
            _heading(doc, s[4:].strip(), 11, bold=True)
            continue
        if s.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, s[2:].strip(), 10)
            continue

        # Zone d'en-tete (avant la 1re section) : titre poste, PRIIME, contact.
        if not header_done:
            if re.fullmatch(r"\*\*[^*]+\*\*", s):     # **Titre du poste**
                _heading(doc, s[2:-2], 13, bold=True)
                continue
            if re.fullmatch(r"\*[^*]+\*", s):         # *Admissible PRIIME...*
                p = doc.add_paragraph(); _add_inline(p, s, 11)
                continue
            if "|" in s:                              # ligne de coordonnees
                p = doc.add_paragraph(); _add_inline(p, s, 10)
                continue

        # Ligne entierement en italique = date (-> droite) ou contexte (-> gauche).
        if re.fullmatch(r"\*[^*]+\*", s):
            inner = s[1:-1]
            p = doc.add_paragraph()
            if _DATE_RE.search(inner) and len(inner) < 45:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(inner); r.italic = True; r.font.name = FONT; r.font.size = Pt(10)
            continue

        p = doc.add_paragraph()
        _add_inline(p, s, 10)

    doc.save(out_path)  # accepte un chemin (str/Path) OU un file-like (BytesIO)


def render_simple(md_text, out_path, font_size=11, margin_cm=2.0):
    """Rendu fidele generique (gere gras ET italique) pour LM / Salaire / Guide."""
    doc, _ = _new_doc(margin_cm=margin_cm)
    for raw in md_text.split("\n"):
        s = raw.strip()
        if s.startswith("# "):
            _heading(doc, s[2:].strip(), 15, bold=True)
        elif s.startswith("## "):
            _heading(doc, s[3:].strip(), 12, bold=True, border=True, upper=True)
        elif s.startswith("### "):
            _heading(doc, s[4:].strip(), 11, bold=True)
        elif s.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet"); _add_inline(p, s[2:].strip(), font_size)
        elif s.startswith("---"):
            doc.add_paragraph()
        elif s:
            p = doc.add_paragraph(); _add_inline(p, s, font_size)
        else:
            doc.add_paragraph()
    doc.save(out_path)  # accepte un chemin (str/Path) OU un file-like (BytesIO)


def render(md_text, out_path, kind):
    kind = (kind or "").lower()
    if kind == "cv":
        render_cv(md_text, out_path)
    elif kind == "lm":
        render_simple(md_text, out_path, font_size=11, margin_cm=2.5)
    else:  # salaire, guide, defaut
        render_simple(md_text, out_path, font_size=11, margin_cm=2.0)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python docx_render.py <input.md> <output.docx> <cv|lm|salaire|guide>", file=sys.stderr)
        sys.exit(1)
    src = Path(sys.argv[1]).read_text(encoding="utf-8")
    render(src, sys.argv[2], sys.argv[3])
    print(f"OK -> {sys.argv[2]}")
